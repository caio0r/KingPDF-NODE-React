from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
import shutil
import os
from pdf2docx import Converter
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from core.utils import cleanup_file

router = APIRouter()

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"

@router.post("/pdf-to-word")
def convert_pdf_to_word(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a PDF.")

    os.makedirs(UPLOAD_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    input_path = os.path.join(UPLOAD_DIR, f"word_in_{file.filename}")
    output_filename = f"{os.path.splitext(file.filename)[0]}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)

    try:
        # Save uploaded file
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Convert using pdf2docx
        cv = Converter(input_path)
        # settings to minimize extra breaks
        settings = {
            'debug': False,
            'margin_bottom': 0,
            'margin_top': 0,
            'margin_left': 0,
            'margin_right': 0,
            'check_font_size': False  # Let Word handle font scaling slightly better?
        }
        cv.convert(output_path, start=0, end=None, **settings)
        cv.close()

        if not os.path.exists(output_path):
             raise Exception("Conversion failed to produce output file")

        # --- AGGRESSIVE POST-PROCESSING ---
        try:
            doc = Document(output_path)
            
            # 1. Extreme Margins (0.5cm)
            # This virtually guarantees content fits on one page if it fit in the PDF.
            for section in doc.sections:
                section.top_margin = Cm(0.5)
                section.bottom_margin = Cm(0.5)
                section.left_margin = Cm(1.5)  # Relaxed for visual balance
                section.right_margin = Cm(1.5) # Relaxed for visual balance
                section.header_distance = Cm(0)
                section.footer_distance = Cm(0)

            # 2. Compact Style Handling
            # Iterate over paragraphs to remove "Space After" which pushes content down.
            for paragraph in doc.paragraphs:
                p_fmt = paragraph.paragraph_format
                # Force single line spacing
                p_fmt.line_spacing = 1.0
                # Remove space before/after paragraph
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)
            
            # 3. Table cleanup (Tables often create overflow)
            for table in doc.tables:
                table.autofit = True
                table.allow_autofit = True
                
            doc.save(output_path)
            print("Aggressive layout cleanup applied.")
            
        except Exception as e:
            print(f"Warning: Layout cleanup failed: {e}")

        # Add background task to clean up the output file after response is sent
        background_tasks.add_task(cleanup_file, output_path)

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=output_filename
        )

    except Exception as e:
        print(f"Error converting PDF to Word: {e}")
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

    finally:
        cleanup_file(input_path)
